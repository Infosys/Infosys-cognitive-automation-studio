'''
Copyright 2020 Infosys Ltd.
Use of this source code is governed by Apache 2.0 license that can be found in the LICENSE file or at 
https://opensource.org/licenses/Apache-2.0 .
'''
from docx import Document
import re
from abstract_bot import Bot

#class for bot
class FindkeywordAndReplaceInWord(Bot):
    #method to initialise 
    def bot_init(self):
        pass
   
    def execute(self,executionContext):
        try:

            documentName =executionContext['documentName']
            keyword=executionContext['keyword']
            replacedWord=executionContext['replacedWord']
            saveAs=executionContext['saveAs']

            if documentName == '':
                return ("Missing argument : documentName")
            if keyword == '':
                return ("Missing argument : keyword")
            if  replacedWord == '':
                return ("Missing argument :replacedWord")
            if  saveAs == '':
                return ("Missing argument :saveAs")

            def docx_replace_regex(doc_obj, regex , replace,match_count):
                for p in doc_obj.paragraphs:
                    if regex.search(p.text):
                        inline = p.runs
            # Loop added to work with runs (strings with same style)
                        for i in range(len(inline)):
                            if regex.search(inline[i].text):
                                match_count+=1
                                text = regex.sub(replace, inline[i].text)
                                #print("Actual Text: ", inline[i].text)
                                #print("Modified Text: ", text)
                                inline[i].text = text
              
                for table in doc_obj.tables:
                    for row in table.rows:
                        #print("row values: ",row.cells)
                        for cell in row.cells:
                            match_count= docx_replace_regex(cell, regex , replace,match_count)
                return(match_count)

            regex1 = re.compile(keyword)
            replace1 = replacedWord
            filename = documentName
            doc = Document(filename)
            final_match_count= docx_replace_regex(doc, regex1 , replace1,0)
            doc.save(saveAs)
            
            if final_match_count==0:
                return {'Status':'No keyword found'}
            else:
                return {'Status':'Success'}
        except Exception as e:
          return {'Exception' : str(e)}
 

if __name__ == "__main__":
    context = {}
    
    bot_obj = FindkeywordAndReplaceInWord()
    
    context = {'documentName':'','keyword':'','replacedWord':'','saveAs':''}
    bot_obj.bot_init()
    #Calling of execute function on linux server
    output = bot_obj.execute(context)
    print(output)  