'''
Copyright 2020 Infosys Ltd.
Use of this source code is governed by Apache 2.0 license that can be found in the LICENSE file or at 
https://opensource.org/licenses/Apache-2.0 .
'''
from spellchecker import SpellChecker
from docx import Document
from abstract_bot import Bot

#class for bot
class AutoCorrectSpellingsInWord(Bot):
    
    def bot_init(self):
        pass
   
    def execute(self,executionContext):
        try:

            sourceFileName=executionContext['sourceFileName']
            destinationFileName=executionContext['destinationFileName']
                     

            # connecting to server
            spell = SpellChecker()
            def correct_spellings(text):
                corrected_text = []
                misspelled_words = spell.unknown(text)
                for word in text:
                    if word in misspelled_words:
                        corrected_text.append(spell.correction(word))
                    else:
                            corrected_text.append(word)
                return " ".join(corrected_text)
            
            doc = Document(sourceFileName)
            resDoc = Document()
            for paragraph in doc.paragraphs:
                p = paragraph.text.split()
                result =correct_spellings(p)
                print(result)
                resDoc.add_paragraph(result)
                resDoc.save(destinationFileName)

            return {'Status':'Success'}
        except Exception as e:
          return {'Exception' : str(e)}
 

if __name__ == "__main__":
    context = {}
    
    bot_obj = AutoCorrectSpellingsInWord()
    
    '''context = {'sourceFileName':'C:\\Users\\moosashah.syed\\Desktop\\snippet.docx',
               'destinationFileName':'C:\\Users\\moosashah.syed\\Desktop\\now.docx'}'''
    
    context = {'sourceFileName':'',
               'destinationFileName':''}
    bot_obj.bot_init()
    #Calling of execute function on linux server
    output = bot_obj.execute(context)
    print(output)  
