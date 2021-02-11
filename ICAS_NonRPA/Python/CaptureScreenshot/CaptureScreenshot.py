'''
Copyright 2020 Infosys Ltd.
Use of this source code is governed by Apache 2.0 license that can be found in the LICENSE file or at 
https://opensource.org/licenses/Apache-2.0 .
'''
from docx import Document
from docx.shared import Inches
import PIL
from PIL import ImageGrab,Image

from abstract_bot import Bot
class CaptureScreenshot(Bot):
    
    def bot_init(self):
        pass
    def execute(self,executionContext):
        try:
            imageName =executionContext['imageName']
            documentName=executionContext['documentName']
            document=Document()
            p=document.add_paragraph()
            r=p.add_run()
            image=ImageGrab.grab()
            wpercent=(1000/float(image.size[0]))
            hpercent=int(float(image.size[1])*float(wpercent))
            image=image.resize((1000,hpercent),PIL.Image.ANTIALIAS)
            image.save(imageName,'JPEG',quality=95)
            r.add_picture(imageName,width=Inches(6),height=Inches(6))
            document.save(documentName)
            
    
            return {'Output':'Word file created successfully'}
        except Exception as e:
            return {'Exception':str(e)}
          
        
        
if __name__ == "__main__":
    context = {}
    #class object creation
    bot_obj = CaptureScreenshot()
    #giving parameter as a dictinoary
    context = {'documentName':'','imageName':''}
    bot_obj.bot_init()
    #Calling of execute function using object of CaptureScreenshot class
    output = bot_obj.execute(context)
    print(output)